from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def muda_texto_documento(doc, campo, valor):
    for p in doc.paragraphs:
        if campo in p.text:
            novo_texto = p.text.replace(campo, valor)
            for run in p.runs:
                run.clear()
            p.clear()
            p.add_run(novo_texto)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if campo in p.text:
                        novo_texto = p.text.replace(campo, valor)
                        for run in p.runs:
                            run.clear()
                        p.clear()
                        p.add_run(novo_texto)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER