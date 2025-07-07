from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana, pega_quantidade_dias_mes
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
# from mysql.connector import Error # Não usado diretamente
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import zipfile
from datetime import datetime, date, time, timedelta
# Imports necessários para pegar_feriados_mes
from dateutil.easter import easter
import holidays
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# import uuid # uuid não está sendo usado

bp_converte_setor_pdf = Blueprint('bp_converte_setor_pdf', __name__)

# Funções a serem adicionadas
def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_row_background(row, color_hex):
    for cell in row.cells:
        set_cell_background(cell, color_hex)


def pegar_feriados_mes(ano, mes, estado='AM'):
    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"

    conexao = connect_mysql()
    cursor = conexao.cursor(dictionary=True)
    feriados_municipais_db = []
    try:
        # A query agora busca a coluna ponto_facultativo
        query_sql = "SELECT data, ponto_facultativo FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        cursor.execute(query_sql, params)
        feriados_municipais_db = cursor.fetchall()
    finally:
        if conexao.is_connected():
            cursor.close()
            conexao.close()

    pontos_facultativos = []
    # feriados_normais = [] # Não é mais necessário
    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        ponto_facultativo = feriado_row.get('ponto_facultativo', 0)
        data_feriado_obj = None
        if data_db is None:
            continue
        if hasattr(data_db, 'date'):
            data_feriado_obj = data_db.date()
        elif isinstance(data_db, date):
            data_feriado_obj = data_db
        else:
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
            except ValueError:
                continue
        if data_feriado_obj:
            if ponto_facultativo:
                pontos_facultativos.append(data_feriado_obj)
            # else: # O feriado normal já é adicionado abaixo
            #     feriados_normais.append(data_feriado_obj)
            br_feriados[data_feriado_obj] = "Feriado Municipal"
            
    feriados_mes = [d for d in br_feriados if d.month == mes]
    pontos_facultativos_mes = [d for d in pontos_facultativos if d.month == mes]
    
    # Retorna as duas listas
    return feriados_mes, pontos_facultativos_mes


def formatar_horario_para_hh_mm_v2(valor_horario):
    if not valor_horario: return ''
    if isinstance(valor_horario, time): return valor_horario.strftime('%H:%M')
    if isinstance(valor_horario, timedelta):
        total_seconds = abs(int(valor_horario.total_seconds()))
        hours = (total_seconds // 3600) % 24
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"
    if isinstance(valor_horario, str):
        try:
            if valor_horario.count(':') == 2: return datetime.strptime(valor_horario, '%H:%M:%S').strftime('%H:%M')
            elif valor_horario.count(':') == 1: return datetime.strptime(valor_horario, '%H:%M').strftime('%H:%M')
            return valor_horario
        except ValueError: return valor_horario
    return str(valor_horario)

@bp_converte_setor_pdf.route('/api/setores/pdf', methods=['POST'])
def converte_setores_pdf():
    conexao_principal = None
    try:
        body = request.json or {}
        setores = body.get('setores')
        mes_body = body.get('mes')

        if not setores or not isinstance(setores, list):
            return jsonify({'erro': 'Nenhum setor selecionado ou formato inválido'}), 400

        arquivos_zip_gerados_todos_setores = []
        mes_por_extenso_geral = data_atual(mes_body)['mes']

        for setor_nome in setores:
            data_ano_mes_atual = data_atual(mes_body)
            mes_numerico = data_ano_mes_atual['mes_numerico']
            ano = data_ano_mes_atual['ano']
            quantidade_dias_no_mes = pega_quantidade_dias_mes(ano, mes_numerico) 

            conexao_principal = connect_mysql()
            cursor = conexao_principal.cursor(dictionary=True)

            query_funcionarios = "SELECT * FROM funcionarios WHERE setor = %s"
            cursor.execute(query_funcionarios, (setor_nome,))
            funcionarios = cursor.fetchall()

            if not funcionarios:
                if conexao_principal and conexao_principal.is_connected():
                    cursor.close()
                    conexao_principal.close()
                print(f"Nenhum funcionário encontrado no setor {setor_nome}, pulando.")
                continue

            arquivos_pdf_gerados_neste_setor = []
            setor_limpo = setor_nome.strip().replace('/', '_')

            for funcionario in funcionarios:
                # --- INÍCIO DO BLOCO CORRIGIDO ---

                # Busca de feriados por funcionário
                estado_funcionario = funcionario.get('estado', 'AM')
                feriados_do_mes, pontos_facultativos_do_mes = pegar_feriados_mes(ano, mes_numerico, estado=estado_funcionario)
                
                # CORREÇÃO 1: A linha de print foi atualizada para usar as variáveis corretas
                print(f"DEBUG: Para {funcionario.get('nome')}, Feriados: {feriados_do_mes}, Pontos Facultativos: {pontos_facultativos_do_mes}")

                # CORREÇÃO 2: Todo o bloco abaixo foi indentado para ficar DENTRO do loop 'for funcionario in funcionarios:'
                template_path = 'FREQUÊNCIA_MENSAL.docx'
                doc = Document(template_path)
                
                cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados_do_mes, pontos_facultativos_do_mes)

                troca_de_dados = {
                    "CAMPO SETOR": funcionario.get('setor', ''),
                    "CAMPO MÊS": mes_por_extenso_geral,
                    "CAMPO NOME": funcionario.get('nome', ''),
                    "CAMPO ANO": str(ano),
                    "CAMPO HORARIO": str(funcionario.get('horario', '')),
                    "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(funcionario.get('horarioentrada', '')),
                    "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(funcionario.get('horariosaida', '')),
                    "CAMPO MATRÍCULA": str(funcionario.get('matricula', '')),
                    "CAMPO CARGO": funcionario.get('cargo', ''),
                }
                for placeholder, valor in troca_de_dados.items():
                    muda_texto_documento(doc, placeholder, valor)

                nome_limpo = funcionario.get('nome', 'NOME_PADRAO').strip().replace('/', '_')
                caminho_pasta = f"setor/{setor_limpo}/{mes_por_extenso_geral}"
                os.makedirs(caminho_pasta, exist_ok=True)
                nome_base = f"FREQUENCIA_{nome_limpo.replace(' ', '_')}"
                docx_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.docx"))
                pdf_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.pdf"))
                doc.save(docx_path)
                convert_to_pdf(docx_path, pdf_path)
                arquivos_pdf_gerados_neste_setor.append(pdf_path)

                cursor.execute(
                    "INSERT INTO arquivos_pdf (servidor_id, caminho_pdf) VALUES (%s, %s)",
                    (funcionario['id'], pdf_path)
                )

                # --- FIM DO BLOCO CORRIGIDO (QUE AGORA ESTÁ DENTRO DO LOOP) ---

            if arquivos_pdf_gerados_neste_setor:
                zip_path_setor = f"setor/{setor_limpo}/frequencias_funcionarios_{setor_limpo}_{mes_por_extenso_geral}.zip"
                with zipfile.ZipFile(zip_path_setor, 'w') as zipf:
                    for pdf in arquivos_pdf_gerados_neste_setor:
                        zipf.write(pdf, os.path.basename(pdf))
                arquivos_zip_gerados_todos_setores.append(zip_path_setor)
                cursor.execute(
                    "INSERT INTO arquivos_zip (setor, mes, caminho_zip, tipo) VALUES (%s, %s, %s, %s)",
                    (setor_nome, mes_por_extenso_geral, zip_path_setor, 'funcionarios_setor')
                )
            
            conexao_principal.commit()
            if conexao_principal and conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()

        # ... (O restante da função para criar o ZIP final continua igual)
        if not arquivos_zip_gerados_todos_setores:
            return jsonify({'message': 'Nenhum arquivo ZIP de funcionários foi gerado.'}), 200

        if len(arquivos_zip_gerados_todos_setores) > 1:
            zip_final_path = f"setor/frequencias_multissetores_funcionarios_{mes_body.replace('/','-')}_{ano}.zip"
            with zipfile.ZipFile(zip_final_path, 'w') as zipf:
                for zip_file in arquivos_zip_gerados_todos_setores:
                    zipf.write(zip_file, os.path.basename(zip_file))
            
            conexao_principal = connect_mysql()
            cursor = conexao_principal.cursor(dictionary=True)
            cursor.execute(
                "INSERT INTO arquivos_zip (setor, mes, ano, caminho_zip, tipo) VALUES (%s, %s, %s, %s, %s)",
                ('multissetores_funcionarios', mes_por_extenso_geral, str(ano), zip_final_path, 'multissetores_funcionarios_geral')
            )
            conexao_principal.commit()
            if conexao_principal and conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()
            return send_file(zip_final_path, mimetype='application/zip', as_attachment=True, download_name=os.path.basename(zip_final_path))
        elif arquivos_zip_gerados_todos_setores:
            return send_file(arquivos_zip_gerados_todos_setores[0], mimetype='application/zip', as_attachment=True, download_name=os.path.basename(arquivos_zip_gerados_todos_setores[0]))
        
        return jsonify({'message': 'Processamento concluído, mas nenhum ZIP para enviar.'}), 200

    except Exception as exception:
        print(f"ERRO ROTA SETOR FUNCIONARIOS: {str(exception)}")
        if conexao_principal and conexao_principal.is_connected():
            if 'cursor' in locals() and cursor:
                cursor.close()
            conexao_principal.close()
        return jsonify({'erro': f'Erro ao processar setores: {str(exception)}'}), 500

# Modificado para aceitar 'feriados' e aplicar a lógica

def cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados, pontos_facultativos):
    """
    Preenche a tabela de frequência mensal no documento Word, ajustando o número de linhas
    e preenchendo os dias com os respectivos status (Sábado, Domingo, Feriado, Ponto Facultativo, Férias).
    """
    linha_inicial = 8

    if not doc.tables:
        print("AVISO: Nenhuma tabela encontrada no documento.")
        return
    
    table = doc.tables[0]

    # 1. Ajustar o número de linhas na tabela para corresponder à quantidade de dias no mês
    target_total_rows_in_table = linha_inicial + quantidade_dias_no_mes

    while len(table.rows) > target_total_rows_in_table:
        row_to_delete = table.rows[-1]
        tbl_element = table._tbl
        tr_element = row_to_delete._tr
        tbl_element.remove(tr_element)

    while len(table.rows) < target_total_rows_in_table:
        new_row = table.add_row()
        # ALTURA DA LINHA AJUSTADA AQUI
        new_row.height = Cm(0.48) 
        new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in new_row.cells:
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 2. Preencher as linhas de dados
    for i in range(quantidade_dias_no_mes):
        dia = i + 1
        row = table.rows[linha_inicial + i]
        
        # APLICA A ALTURA AJUSTADA TAMBÉM NAS LINHAS EXISTENTES PARA GARANTIR CONSISTÊNCIA
        row.height = Cm(0.48)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        data_atual_obj = date(ano, mes_numerico, dia)
        dia_semana = pega_final_de_semana(ano, mes_numerico, dia)

        # Limpeza e formatação inicial da linha
        for cell in row.cells:
            cell.text = "" 
            for paragraph in cell.paragraphs:
                paragraph.clear() 
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Preencher o número do dia na primeira coluna
        dia_cell = row.cells[0]
        dia_paragraph = dia_cell.paragraphs[0]
        dia_run = dia_paragraph.add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)
        dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Lógica para Sábados e Domingos
        texto_status = ""
        if dia_semana == 5:
            texto_status = "SÁBADO"
        elif dia_semana == 6:
            texto_status = "DOMINGO"

        if texto_status:
            set_row_background(row, 'C5E0B4') 
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = texto_status
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(6) 

        # Lógica para Ponto Facultativo e Feriado (apenas em dias de semana)
        if data_atual_obj in pontos_facultativos and dia_semana not in [5, 6]:
            set_row_background(row, 'C5E0B4')
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = "PONTO FACULTATIVO"
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(6)
        
        elif data_atual_obj in feriados and dia_semana not in [5, 6]:
            set_row_background(row, 'C5E0B4')
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = "FERIADO"
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(6)

        # Lógica para Férias (sobrescreve o status anterior se for o caso, exceto em fins de semana)
        if funcionario.get('feriasinicio') and funcionario.get('feriasfinal'):
            ferias_inicio_raw = funcionario['feriasinicio']
            ferias_final_raw = funcionario['feriasfinal']
            ferias_inicio = ferias_inicio_raw.date() if hasattr(ferias_inicio_raw, 'date') else ferias_inicio_raw
            ferias_final = ferias_final_raw.date() if hasattr(ferias_final_raw, 'date') else ferias_final_raw

            if isinstance(ferias_inicio, date) and isinstance(ferias_final, date) and \
               (ferias_inicio <= data_atual_obj <= ferias_final and dia_semana not in [5, 6]):
                
                set_row_background(row, 'C5E0B4')
                for j in [2, 5, 9, 13]:
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = "FÉRIAS"
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.name = "Calibri"
                                run.font.size = Pt(6)
    
    # Se você tem múltiplas tabelas no documento e só quer processar a primeira,
    # o loop `for table in doc.tables:` pode ser removido ou adicionar um `break` no final.
    # Se o código original processava todas, mantenha o loop. Pela sua estrutura, parece que
    # a intenção é processar a primeira tabela de frequência encontrada.
    # Adicionando um break para processar apenas a primeira tabela, que é o comportamento mais comum.
    # Remova este 'break' se você intencionalmente processa múltiplas tabelas de frequência no mesmo doc.
    # No seu código original, não havia break, então o loop for table continuaria.
    # Para este ajuste de linhas, faz mais sentido focar em UMA tabela principal.
    # Se o loop for table for mantido, a lógica de ajuste de linhas será aplicada a cada tabela.
