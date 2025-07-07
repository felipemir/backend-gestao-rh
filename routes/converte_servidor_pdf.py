from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana, pega_quantidade_dias_mes
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
from mysql.connector import Error
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import holidays
from datetime import datetime, date, timedelta,time
import zipfile
import re
from datetime import date
from dateutil.easter import easter
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


bp_converte_servidor_pdf = Blueprint('bp_converte_servidor_pdf', __name__)



def set_cell_background(cell, color_hex):
    """
    Define a cor de fundo da célula (color_hex no formato 'RRGGBB', ex: 'B7DEE8' para azul claro).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_row_background(row, color_hex):
    for cell in row.cells:
        set_cell_background(cell, color_hex)
    
# ...existing code...
def pegar_feriados_mes(ano, mes, estado='AM'):
    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"

    conexao = connect_mysql()
    cursor = conexao.cursor(dictionary=True)
    feriados_municipais_db = []
    try:
        query_sql = "SELECT data, ponto_facultativo FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        cursor.execute(query_sql, params)
        feriados_municipais_db = cursor.fetchall()
    finally:
        if conexao.is_connected():
            cursor.close()
            conexao.close()

    pontos_facultativos = []
    feriados_normais = []
    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        ponto_facultativo = feriado_row.get('ponto_facultativo', 0)
        data_feriado_obj = None
        if data_db is None:
            continue
        if hasattr(data_db, 'date'):
            data_feriado_obj = data_db.date()
        elif isinstance(data_db, date):
            data_feriado_obj = data_db
        else:
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
            except ValueError:
                continue
        if data_feriado_obj:
            if ponto_facultativo:
                pontos_facultativos.append(data_feriado_obj)
            else:
                feriados_normais.append(data_feriado_obj)
            br_feriados[data_feriado_obj] = "Feriado Municipal"
    feriados_mes = [d for d in br_feriados if d.month == mes]
    pontos_facultativos_mes = [d for d in pontos_facultativos if d.month == mes]
    return feriados_mes, pontos_facultativos_mes

def limpa_nome(nome):
    return re.sub(r'[^\w\s-]', '', nome).strip().replace(' ', '_')


def formatar_horario_para_hh_mm_v2(valor_horario):
    """
    Formata um valor de horário para o formato HH:MM, removendo os segundos.
    """
    if not valor_horario:  # Se for None, string vazia, etc.
        return ''

    # Caso 1: Se for um objeto datetime.time
    if isinstance(valor_horario, time):
        return valor_horario.strftime('%H:%M')

    # Caso 2: Se for um objeto datetime.timedelta (comum de bancos de dados para colunas TIME)
    if isinstance(valor_horario, timedelta):
        total_seconds = int(valor_horario.total_seconds())
        # Ignora dias, foca apenas na parte de tempo do dia
        if total_seconds < 0: # Lida com timedeltas negativos, se aplicável
            # Você pode querer um tratamento específico aqui, por ex., '' ou erro
            # Para simplificar, vamos assumir horas e minutos a partir de 0 se for negativo
            # ou tratar como 00:00. A lógica exata pode depender do seu caso de uso.
            # Exemplo: tratar como 00:00 se negativo ou converter para positivo
            # Para este exemplo, vamos apenas calcular com base no valor absoluto.
            total_seconds = abs(total_seconds)

        hours = (total_seconds // 3600) % 24 # Garante que as horas fiquem dentro de 0-23
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"

    # Caso 3: Se for uma string
    if isinstance(valor_horario, str):
        try:
            # Tenta primeiro como HH:MM:SS
            if valor_horario.count(':') == 2:
                dt_obj = datetime.strptime(valor_horario, '%H:%M:%S')
                return dt_obj.strftime('%H:%M')
            # Depois como HH:MM
            elif valor_horario.count(':') == 1:
                dt_obj = datetime.strptime(valor_horario, '%H:%M')
                return dt_obj.strftime('%H:%M') # Já está no formato, mas re-formata para garantir
            else:
                # Se não for um formato de tempo reconhecido, retorna a string original
                return valor_horario
        except ValueError:
            # Se a conversão da string falhar
            return valor_horario # Retorna a string original

    # Fallback: Se não for nenhum dos tipos acima, tenta converter para string
    return str(valor_horario)

@bp_converte_servidor_pdf.route('/api/servidores/pdf', methods=['POST'])
def converte_servidor_pdf():
    try:
        body = request.json or {}
        funcionarios_id = body.get('funcionarios', [])

        if not funcionarios_id:
            return jsonify({'erro': 'Nenhum funcionário selecionado'}), 400

        try:
            ids = [int(id) for id in funcionarios_id]
        except ValueError:
            return jsonify({'erro': 'IDs inválidos'}), 400

        mes_body = body.get('mes')
        data_ano_mes_atual = data_atual(mes_body)
        mes_por_extenso = data_ano_mes_atual['mes']
        mes_numerico = data_ano_mes_atual['mes_numerico']
        ano = data_ano_mes_atual['ano']
        quantidade_dias_no_mes = pega_quantidade_dias_mes(ano, mes_numerico)

        conexao = connect_mysql()
        cursor = conexao.cursor(dictionary=True)

        placeholders = ','.join(['%s'] * len(ids))
        query = f"SELECT * FROM funcionarios WHERE id IN ({placeholders})"
        cursor.execute(query, ids)
        funcionarios = cursor.fetchall()

        if not funcionarios:
            conexao.close()
            print(" cair aquiiii")
            return jsonify({'erro': 'Nenhum funcionário encontrado'}), 404

        arquivos_gerados = []
    
        for funcionario in funcionarios:
            estado_funcionarios = funcionario.get('estado', 'AM')
            feriados_do_mes, pontos_facultativos_mes = pegar_feriados_mes(ano, mes_numerico, estado=estado_funcionarios)
            template_path = 'FREQUÊNCIA_MENSAL.docx'
            doc = Document(template_path)

            cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados_do_mes, pontos_facultativos_mes)
            # Formatar horário
            troca_de_dados = {
            "CAMPO SETOR": funcionario['setor'],
            "CAMPO MÊS": mes_por_extenso,
            "CAMPO NOME": funcionario['nome'],
            "CAMPO ANO": str(ano),
            "CAMPO HORARIO": funcionario.get('horario', ''), # Mantido como está, sem formatação específica aqui
            "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(funcionario.get('horarioentrada', '')),
            "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(funcionario.get('horariosaida', '')),
            "CAMPO MATRÍCULA": str(funcionario.get('matricula', '')),
            "CAMPO CARGO": funcionario.get('cargo', ''),
        }

            for placeholder, valor in troca_de_dados.items():
                muda_texto_documento(doc, placeholder, valor)

            nome_limpo = limpa_nome(funcionario['nome'])
            setor_limpo = limpa_nome(funcionario['setor'])
            caminho_pasta = f"setor/{setor_limpo}/servidor/{mes_por_extenso}/{nome_limpo}"
            os.makedirs(caminho_pasta, exist_ok=True)

            nome_base = f"{nome_limpo}_FREQUENCIA"
            docx_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.docx"))
            pdf_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.pdf"))

            doc.save(docx_path)
            convert_to_pdf(docx_path, pdf_path)

            arquivos_gerados.append(pdf_path)

        # Criar ZIP com todos os PDFs
        zip_path = os.path.abspath(f"setor/frequencias_{mes_por_extenso}.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for pdf in arquivos_gerados:
                zipf.write(pdf, os.path.basename(pdf))

        # Salvar caminho do ZIP no banco
        cursor.execute(
            "INSERT INTO arquivos_zip (mes, caminho_zip, tipo) VALUES (%s, %s, %s)",
            (mes_por_extenso, zip_path, 'servidores')
        )

        conexao.commit()
        conexao.close()

        return send_file(
            zip_path,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'frequencias_servidores_{mes_por_extenso}.zip'
        )

    except Exception as exception:
        if 'conexao' in locals():
            conexao.close()
        return jsonify({'erro': f'Erro: {str(exception)}'}), 500

def cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados,pontos_facultativos):
    linha_inicial = 8

    if not doc.tables:
        print("AVISO: Nenhum tabela encontrada no documento.")
        return
    
    table = doc.tables[0] # Assume-se que a primeira tabela é a de frequência

    # 1. Aplicar formatação base em todas as linhas existentes (como no seu código original)
    # Esta formatação pode ser muito genérica; idealmente, o template já teria os estilos corretos
    # para cabeçalhos vs. dados, mas vamos manter sua lógica original por enquanto.
    for row in table.rows:
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7)
                    run.font.bold = False

    # 2. Ajustar o número de linhas na tabela para corresponder à quantidade_dias_no_mes
    # Linhas de dados necessárias = quantidade_dias_no_mes
    # Total de linhas que a tabela deve ter = linha_inicial (cabeçalho) + quantidade_dias_no_mes
    target_total_rows_in_table = linha_inicial + quantidade_dias_no_mes

    # Remover linhas excedentes do final da tabela
    while len(table.rows) > target_total_rows_in_table:
        row_to_delete = table.rows[-1] # Pega a última linha da tabela
        tbl_element = table._tbl
        tr_element = row_to_delete._tr
        tbl_element.remove(tr_element)
        print(f"INFO: Linha excedente removida. Total de linhas agora: {len(table.rows)}")

    # Adicionar linhas se estiverem faltando
    while len(table.rows) < target_total_rows_in_table:
        new_row = table.add_row()
        new_row.height = Cm(0.5) # Aplicar altura padrão às novas linhas
        new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Garante que as células da nova linha tenham parágrafos formatados
        for cell in new_row.cells:
            # Assegura que existe pelo menos um parágrafo e o alinha
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Pode-se definir a fonte padrão para o parágrafo ou para um run vazio aqui, se necessário
            # Mas geralmente o estilo da tabela ou a formatação do conteúdo adicionado depois cuidará disso.
        print(f"INFO: Linha faltante adicionada. Total de linhas agora: {len(table.rows)}")

    # 3. Preencher as linhas de dados (seu código original a partir daqui)
    # Defina pontos_facultativos antes do loop, por exemplo, como uma lista vazia ou conforme sua lógica
    #pontos_facultativos_mes = []  # Substitua por sua lógica para obter pontos facultativos, se necessário

    for i in range(quantidade_dias_no_mes):
        dia = i + 1
        # Agora é seguro acessar table.rows[linha_inicial + i]
        row = table.rows[linha_inicial + i]
        data_atual = date(ano, mes_numerico, dia) # Use o nome data_atual como no seu código
        dia_semana = pega_final_de_semana(ano, mes_numerico, dia) # Assume que esta função existe

        # Limpeza das células da linha atual antes de preencher
        for cell in row.cells:
            cell.text = "" # Limpa o conteúdo principal da célula (primeiro parágrafo)
            for paragraph in cell.paragraphs: # Itera sobre todos os parágrafos
                paragraph.clear() # Limpa todos os 'runs' (texto formatado) de cada parágrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Garante o alinhamento

        # Preencher número do dia
        dia_cell = row.cells[0]
        # Garante que há um parágrafo para adicionar o run
        dia_paragraph = dia_cell.paragraphs[0] if dia_cell.paragraphs else dia_cell.add_paragraph()
        dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Reafirma o alinhamento se for um novo parágrafo
        dia_run = dia_paragraph.add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)
        # dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Já definido

        # Lógica para Sábados, Domingos, Feriados e Férias
        # A ordem das suas verificações originais define a prioridade (a última condição que sobrescreve cell.text vence)
        
        texto_status = "" # Para Sábado/Domingo

        if dia_semana == 5:
            texto_status = "SÁBADO"
        elif dia_semana == 6:
            texto_status = "DOMINGO"

        if texto_status: # Escreve SÁBADO ou DOMINGO
            set_row_background(row, 'C5E0B4') # VERDE
            for j in [2, 5, 9, 13]: # Seus índices de coluna originais
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = texto_status # Define o texto, limpando parágrafos anteriores
                    # Reaplicar formatação após cell.text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs: # O texto agora está em um ou mais runs
                            run.font.bold = True
                            run.font.name = "Calibri" # Garantir consistência
                            run.font.size = Pt(7)     # Garantir consistência
                else:
                    print(f"AVISO: Índice de coluna {j} para S/D fora dos limites.")

        # Corrigido: bloco corretamente indentado e pontos_facultativos definido
        if data_atual in pontos_facultativos and dia_semana not in [5, 6]:
            set_row_background(row, 'C5E0B4')  # 
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = "PONTO FACULTATIVO"
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(7)
                else:
                    print(f"AVISO: Índice de coluna {j} para PONTO FACULTATIVO fora dos limites.")

        # Fer (exceto se for sábado ou domingo) - sobrescreve células se for o caso
        elif data_atual in feriados and dia_semana not in [5, 6]:
            set_row_background(row, 'C5E0B4') # VERDE
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = "FERIADO"

                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(7)
                else:
                    print(f"AVISO: Índice de coluna {j} para FERIADO fora dos limites.")


        # Férias (exceto fins de semana) - sobrescreve células se for o caso
        if funcionario.get('feriasinicio') and funcionario.get('feriasfinal'):
            ferias_inicio_raw = funcionario['feriasinicio']
            ferias_final_raw = funcionario['feriasfinal']
            ferias_inicio = ferias_inicio_raw.date() if hasattr(ferias_inicio_raw, 'date') else ferias_inicio_raw
            ferias_final = ferias_final_raw.date() if hasattr(ferias_final_raw, 'date') else ferias_final_raw

            if isinstance(ferias_inicio, date) and isinstance(ferias_final, date) and \
               (ferias_inicio <= data_atual <= ferias_final and dia_semana not in [5, 6]):
                set_row_background(row, 'C5E0B4') # VERDE   
                for j in [2, 5, 9, 13]:
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = "FÉRIAS"
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.name = "Calibri"
                                run.font.size = Pt(7)
                    else:
                        print(f"AVISO: Índice de coluna {j} para FÉRIAS fora dos limites.")
