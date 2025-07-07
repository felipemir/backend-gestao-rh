from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import zipfile
from datetime import datetime, timedelta, date, time
from dateutil.easter import easter
import holidays
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

bp_converte_setor_estagiario_pdf = Blueprint('bp_converte_setor_estagiario_pdf', __name__)

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_row_background(row, color_hex):
    for cell in row.cells:
        set_cell_background(cell, color_hex)

def pegar_feriados_mes(ano, mes, estado='AM'):
    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"
    conexao = connect_mysql()
    cursor = conexao.cursor(dictionary=True)
    feriados_municipais_db = []
    try:
        query_sql = "SELECT data, ponto_facultativo FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        cursor.execute(query_sql, params)
        feriados_municipais_db = cursor.fetchall()
    finally:
        if conexao.is_connected():
            cursor.close()
            conexao.close()
    pontos_facultativos = []
    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        ponto_facultativo = feriado_row.get('ponto_facultativo', 0)
        data_feriado_obj = None
        if data_db is None:
            continue
        if hasattr(data_db, 'date'):
            data_feriado_obj = data_db.date()
        elif isinstance(data_db, date):
            data_feriado_obj = data_db
        else:
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
            except ValueError:
                continue
        if data_feriado_obj:
            if ponto_facultativo:
                pontos_facultativos.append(data_feriado_obj)
            br_feriados[data_feriado_obj] = "Feriado Municipal"
    feriados_mes = [d for d in br_feriados if d.month == mes]
    pontos_facultativos_mes = [d for d in pontos_facultativos if d.month == mes]
    return feriados_mes, pontos_facultativos_mes

def formatar_horario_para_hh_mm_v2(valor_horario):
    if not valor_horario: return ''
    if isinstance(valor_horario, time): return valor_horario.strftime('%H:%M')
    if isinstance(valor_horario, timedelta):
        total_seconds = abs(int(valor_horario.total_seconds()))
        hours = (total_seconds // 3600) % 24
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"
    if isinstance(valor_horario, str):
        try:
            if valor_horario.count(':') == 2: return datetime.strptime(valor_horario, '%H:%M:%S').strftime('%H:%M')
            elif valor_horario.count(':') == 1: return datetime.strptime(valor_horario, '%H:%M').strftime('%H:%M')
            return valor_horario
        except ValueError: return valor_horario
    return str(valor_horario)

@bp_converte_setor_estagiario_pdf.route('/api/setores/estagiar/pdf', methods=['POST'])
def converte_setores_estagiarios_pdf():
    conexao_principal = None
    try:
        body = request.json or {}
        setores_nomes = body.get('setores')
        mes_body = body.get('mes')

        if not setores_nomes or not isinstance(setores_nomes, list):
            return jsonify({'erro': 'Nenhum setor selecionado ou formato inválido'}), 400

        data_ano_mes_atual = data_atual(mes_body)
        mes_por_extenso = data_ano_mes_atual['mes']
        mes_numerico = data_ano_mes_atual['mes_numerico']
        ano = data_ano_mes_atual['ano']

        # --- LÓGICA DE BUSCA DE FERIADOS LIMPA E CORRIGIDA ---
        estado_para_feriados = 'AM'
        feriados_corrente, pontos_fac_corrente = pegar_feriados_mes(ano, mes_numerico, estado=estado_para_feriados)
        ano_proximo = ano
        mes_proximo = mes_numerico + 1
        if mes_proximo > 12:
            mes_proximo = 1
            ano_proximo += 1
        feriados_proximo, pontos_fac_proximo = pegar_feriados_mes(ano_proximo, mes_proximo, estado=estado_para_feriados)
        
        todos_feriados_do_periodo = list(set(feriados_corrente + feriados_proximo))
        todos_pontos_facultativos_do_periodo = list(set(pontos_fac_corrente + pontos_fac_proximo))
        # --- FIM DA LÓGICA DE BUSCA DE FERIADOS ---

        arquivos_zip_dos_setores = []

        for setor_nome in setores_nomes:
            conexao_principal = connect_mysql()
            cursor = conexao_principal.cursor(dictionary=True)
            query = "SELECT * FROM estagiarios WHERE setor = %s"
            cursor.execute(query, (setor_nome,))
            estagiarios = cursor.fetchall()

            if not estagiarios:
                if conexao_principal and conexao_principal.is_connected():
                    cursor.close()
                    conexao_principal.close()
                print(f"Nenhum estagiário encontrado no setor {setor_nome}, pulando.")
                continue

            arquivos_pdf_gerados_neste_setor = []
            setor_limpo = setor_nome.strip().replace('/', '_')
            caminho_pasta_base_setor = f"setor/estagiarios/{setor_limpo}/{mes_por_extenso}"
            os.makedirs(caminho_pasta_base_setor, exist_ok=True)

            for estagiario in estagiarios:
                template_path = 'FREQUÊNCIA ESTAGIÁRIOS - MODELO.docx'
                doc = Document(template_path)
                
                # Passando as duas listas corretamente
                cria_dias_da_celula(doc, ano, mes_numerico, estagiario, todos_feriados_do_periodo, todos_pontos_facultativos_do_periodo)

                troca_de_dados = {
                    "CAMPO SETOR": estagiario.get('setor', ''),
                    "CAMPO MES": mes_por_extenso,
                    "CAMPO NOME": estagiario.get('nome', ''),
                    "CAMPO ANO": str(ano),
                    "CAMPO HORARIO": str(estagiario.get('horario', '')),
                    "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_entrada', '')),
                    "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_saida', '')),
                    "CAMPO CARGO": str(estagiario.get('cargo', '')),
                }
                for placeholder, valor in troca_de_dados.items():
                    muda_texto_documento(doc, placeholder, valor)

                nome_limpo = estagiario.get('nome', 'NOME_PADRAO').strip().replace('/', '_')
                nome_base = f"FREQUENCIA_ESTAGIARIO_{nome_limpo.replace(' ', '_')}"
                docx_path = os.path.abspath(os.path.join(caminho_pasta_base_setor, f"{nome_base}.docx"))
                pdf_path = os.path.abspath(os.path.join(caminho_pasta_base_setor, f"{nome_base}.pdf"))
                doc.save(docx_path)
                convert_to_pdf(docx_path, pdf_path)
                arquivos_pdf_gerados_neste_setor.append(pdf_path)
                
                cursor.execute("INSERT INTO arquivos_pdf (servidor_id, caminho_pdf) VALUES (%s, %s)", (estagiario['id'], pdf_path))

            if arquivos_pdf_gerados_neste_setor:
                zip_path_setor = f"setor/estagiarios/{setor_limpo}/frequencias_estagiarios_{setor_limpo}_{mes_por_extenso}.zip"
                with zipfile.ZipFile(zip_path_setor, 'w') as zipf:
                    for pdf in arquivos_pdf_gerados_neste_setor:
                        zipf.write(pdf, os.path.basename(pdf))
                arquivos_zip_dos_setores.append(zip_path_setor)
                cursor.execute("INSERT INTO arquivos_zip (setor, mes, caminho_zip, tipo) VALUES (%s, %s, %s, %s)", (setor_nome, mes_por_extenso, zip_path_setor, 'estagiarios_setor'))
            
            conexao_principal.commit()
            if conexao_principal and conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()
        
        if not arquivos_zip_dos_setores:
            return jsonify({'message': 'Nenhum arquivo ZIP de estagiários foi gerado.'}), 200

        if len(arquivos_zip_dos_setores) > 1:
            zip_final_path = f"setor/estagiarios/frequencias_multissetores_estagiarios_{mes_body.replace('/','-')}_{ano}.zip"
            with zipfile.ZipFile(zip_final_path, 'w') as zipf:
                for zip_file_setor in arquivos_zip_dos_setores:
                    zipf.write(zip_file_setor, os.path.basename(zip_file_setor))
            conexao_principal = connect_mysql()
            cursor = conexao_principal.cursor(dictionary=True)
            cursor.execute("INSERT INTO arquivos_zip (setor, mes, ano, caminho_zip, tipo) VALUES (%s, %s, %s, %s, %s)", ('multiestagiarios', mes_por_extenso, str(ano), zip_final_path, 'multiestagiarios_geral'))
            conexao_principal.commit()
            if conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()
            return send_file(zip_final_path, mimetype='application/zip', as_attachment=True, download_name=os.path.basename(zip_final_path))
        elif arquivos_zip_dos_setores:
            return send_file(arquivos_zip_dos_setores[0], mimetype='application/zip', as_attachment=True, download_name=os.path.basename(arquivos_zip_dos_setores[0]))
        
        return jsonify({'message': 'Processamento concluído, mas nenhum ZIP para enviar.'}), 200

    except Exception as exception:
        print(f"ERRO ROTA SETOR ESTAGIARIOS: {str(exception)}")
        if conexao_principal and conexao_principal.is_connected():
            if 'cursor' in locals() and cursor:
                cursor.close()
            conexao_principal.close()
        return jsonify({'erro': f'Erro ao processar setores de estagiários: {str(exception)}'}), 500

def cria_dias_da_celula(doc, ano_param, mes_param, estagiario, feriados, pontos_facultativos):
    """
    Preenche a tabela de frequência de estagiário (período de 21 a 20),
    ajustando linhas e preenchendo os dias com os respectivos status.
    """
    def calcula_periodo_21_a_20(ano_calc, mes_calc):
        data_inicio = datetime(ano_calc, mes_calc, 21)
        if mes_calc == 12:
            data_fim = datetime(ano_calc + 1, 1, 20)
        else:
            data_fim = datetime(ano_calc, mes_calc + 1, 20)
        
        dias_periodo = []
        data_atual_loop = data_inicio
        while data_atual_loop <= data_fim:
            dias_periodo.append({
                "dia": data_atual_loop.day,
                "mes": data_atual_loop.month,
                "ano": data_atual_loop.year
            })
            data_atual_loop += timedelta(days=1)
        return dias_periodo

    linha_inicial = 7
    if not doc.tables:
        print("AVISO: Nenhuma tabela encontrada no documento de estagiário.")
        return
        
    table = doc.tables[0]
    
    # Ajusta o número de linhas na tabela para o período
    dias_periodo = calcula_periodo_21_a_20(ano_param, mes_param)
    target_total_rows = linha_inicial + len(dias_periodo)

    while len(table.rows) > target_total_rows:
        row_to_delete = table.rows[-1]
        tbl_element = table._tbl
        tr_element = row_to_delete._tr
        tbl_element.remove(tr_element)

    while len(table.rows) < target_total_rows:
        new_row = table.add_row()
        # ALTURA DA LINHA AJUSTADA AQUI
        new_row.height = Cm(0.5) 
        new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in new_row.cells:
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Preenche os dados de cada dia do período
    for i, dia_info in enumerate(dias_periodo):
        dia = dia_info["dia"]
        mes_iter = dia_info["mes"]
        ano_dia = dia_info["ano"]
        
        row = table.rows[linha_inicial + i]
        
        # APLICA A ALTURA AJUSTADA TAMBÉM NAS LINHAS EXISTENTES PARA GARANTIR CONSISTÊNCIA
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        current_date_obj = date(ano_dia, mes_iter, dia)
        dia_semana = pega_final_de_semana(ano_dia, mes_iter, dia)
        
        # Limpeza da linha
        for cell in row.cells:
            cell.text = ""
            for paragraph in cell.paragraphs:
                paragraph.clear()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Preenche o dia
        dia_run = row.cells[0].paragraphs[0].add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)

        # Lógica para definir o status do dia
        text_to_write = None
        
        if dia_semana == 5:
            text_to_write = "SÁBADO"
        elif dia_semana == 6:
            text_to_write = "DOMINGO"

        if text_to_write is None:
            if estagiario.get('feriasinicio') and estagiario.get('feriasfinal'):
                ferias_inicio_raw = estagiario['feriasinicio']
                ferias_final_raw = estagiario['feriasfinal']
                ferias_inicio = ferias_inicio_raw.date() if hasattr(ferias_inicio_raw, 'date') else ferias_inicio_raw
                ferias_final = ferias_final_raw.date() if hasattr(ferias_final_raw, 'date') else ferias_final_raw
                if isinstance(ferias_inicio, date) and isinstance(ferias_final, date) and (ferias_inicio <= current_date_obj <= ferias_final):
                    text_to_write = "FÉRIAS"
            
            if text_to_write is None and current_date_obj in pontos_facultativos:
                text_to_write = "PONTO FACULTATIVO"
            
            elif text_to_write is None and current_date_obj in feriados:
                text_to_write = "FERIADO"

        # Se algum status foi definido, escreve nas colunas
        if text_to_write:
            set_row_background(row, 'C5E0B4')
            column_indices_to_fill = [2, 5, 9, 13]

            for j_idx in column_indices_to_fill:
                if j_idx < len(row.cells):
                    cell_to_mark = row.cells[j_idx]
                    p_to_mark = cell_to_mark.paragraphs[0]
                    run_to_mark = p_to_mark.add_run(text_to_write)
                    run_to_mark.font.bold = True
                    run_to_mark.font.name = "Calibri"
                    run_to_mark.font.size = Pt(6)
                    p_to_mark.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
